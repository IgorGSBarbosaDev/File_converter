import os
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
import logging

# Configurar o registro de logs
logging.basicConfig(filename="conversion.log", level=logging.INFO, format="%(asctime)s - %(message)s")


class FileConverter:
    """Lida com a lógica de conversão de arquivos."""

    conversion_functions = {
        "PDF para DOCX": (lambda pdf, docx: FileConverter.convert_pdf_to_docx(pdf, docx), ".docx"),
        "DOCX para PDF": (lambda docx, pdf: FileConverter.convert_docx_to_pdf(docx, pdf), ".pdf"),
        "DOCX para TXT": (lambda docx, txt: FileConverter.convert_docx_to_txt(docx, txt), ".txt"),
        "TXT para DOCX": (lambda txt, docx: FileConverter.convert_txt_to_docx(txt, docx), ".docx"),
        "XLS para CSV": (lambda xls, csv: FileConverter.convert_xls_to_csv(xls, csv), ".csv"),
        "CSV para XLS": (lambda csv, xls: FileConverter.convert_csv_to_xls(csv, xls), ".xls"),
    }

    @staticmethod
    def convert_pdf_to_docx(pdf_file, docx_file):
        try:
            reader = PdfReader(pdf_file)
            doc = Document()
            for page in reader.pages:
                doc.add_paragraph(page.extract_text())
            doc.save(docx_file)
        except Exception as e:
            logging.error(f"Erro ao converter PDF para DOCX: {str(e)}")
            raise

    @staticmethod
    def convert_docx_to_pdf(docx_file, pdf_file):
        try:
            os.system(f'libreoffice --headless --convert-to pdf "{docx_file}" --outdir "{os.path.dirname(pdf_file)}"')
        except Exception as e:
            logging.error(f"Erro ao converter DOCX para PDF: {str(e)}")
            raise

    @staticmethod
    def convert_docx_to_txt(docx_file, txt_file):
        try:
            doc = Document(docx_file)
            with open(txt_file, 'w', encoding='utf-8') as f:
                for para in doc.paragraphs:
                    f.write(para.text + '\n')
        except Exception as e:
            logging.error(f"Erro ao converter DOCX para TXT: {str(e)}")
            raise

    @staticmethod
    def convert_txt_to_docx(txt_file, docx_file):
        try:
            doc = Document()
            with open(txt_file, 'r', encoding='utf-8') as f:
                for line in f:
                    doc.add_paragraph(line.strip())
            doc.save(docx_file)
        except Exception as e:
            logging.error(f"Erro ao converter TXT para DOCX: {str(e)}")
            raise

    @staticmethod
    def convert_xls_to_csv(xls_file, csv_file):
        try:
            df = pd.read_excel(xls_file)
            df.to_csv(csv_file, index=False)
        except Exception as e:
            logging.error(f"Erro ao converter XLS para CSV: {str(e)}")
            raise

    @staticmethod
    def convert_csv_to_xls(csv_file, xls_file):
        try:
            df = pd.read_csv(csv_file)
            df.to_excel(xls_file, index=False)
        except Exception as e:
            logging.error(f"Erro ao converter CSV para XLS: {str(e)}")
            raise