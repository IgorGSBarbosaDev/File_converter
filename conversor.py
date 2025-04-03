import os
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import logging

# Configure logging
logging.basicConfig(filename="conversion.log", level=logging.INFO, format="%(asctime)s - %(message)s")


class FileConverter:
    """LÓGICA DA CONVERSÃO DE ARQUIVOS."""

    conversion_functions = {
        "PDF para DOCX": (lambda pdf, docx: FileConverter.convert_pdf_to_docx(pdf, docx), ".docx"),
        "DOCX para PDF": (lambda docx, pdf: FileConverter.convert_docx_to_pdf(docx, pdf), ".pdf"),
        "DOCX para TXT": (lambda docx, txt: FileConverter.convert_docx_to_txt(docx, txt), ".txt"),
        "TXT para DOCX": (lambda txt, docx: FileConverter.convert_txt_to_docx(txt, docx), ".docx"),
        "XLS para CSV": (lambda xls, csv: FileConverter.convert_xls_to_csv(xls, csv), ".csv"),
        "CSV para XLS": (lambda csv, xls: FileConverter.convert_csv_to_xls(csv, xls), ".xls"),
    }

    @staticmethod
    def convert_pdf_to_docx(pdf_file, docx_file):
        try:
            reader = PdfReader(pdf_file)
            doc = Document()
            for page in reader.pages:
                doc.add_paragraph(page.extract_text())
            doc.save(docx_file)
        except Exception as e:
            logging.error(f"Error converting PDF to DOCX: {str(e)}")
            raise

    @staticmethod
    def convert_docx_to_pdf(docx_file, pdf_file):
        try:
            os.system(f'libreoffice --headless --convert-to pdf "{docx_file}" --outdir "{os.path.dirname(pdf_file)}"')
        except Exception as e:
            logging.error(f"Error converting DOCX to PDF: {str(e)}")
            raise

    @staticmethod
    def convert_docx_to_txt(docx_file, txt_file):
        try:
            doc = Document(docx_file)
            with open(txt_file, 'w', encoding='utf-8') as f:
                for para in doc.paragraphs:
                    f.write(para.text + '\n')
        except Exception as e:
            logging.error(f"Error converting DOCX to TXT: {str(e)}")
            raise

    @staticmethod
    def convert_txt_to_docx(txt_file, docx_file):
        try:
            doc = Document()
            with open(txt_file, 'r', encoding='utf-8') as f:
                for line in f:
                    doc.add_paragraph(line.strip())
            doc.save(docx_file)
        except Exception as e:
            logging.error(f"Error converting TXT to DOCX: {str(e)}")
            raise

    @staticmethod
    def convert_xls_to_csv(xls_file, csv_file):
        try:
            df = pd.read_excel(xls_file)
            df.to_csv(csv_file, index=False)
        except Exception as e:
            logging.error(f"Error converting XLS to CSV: {str(e)}")
            raise

    @staticmethod
    def convert_csv_to_xls(csv_file, xls_file):
        try:
            df = pd.read_csv(csv_file)
            df.to_excel(xls_file, index=False)
        except Exception as e:
            logging.error(f"Error converting CSV to XLS: {str(e)}")
            raise


class FileConverterApp:
    """ CONEXÃO COM MENU E INTERFACE GRÁFICA."""

    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Conversor de Arquivos")
        self.root.configure(bg="#f0f0f0")

        self.input_file_path = tk.StringVar()
        self.conversion_type_var = tk.StringVar()

        self.setup_ui()

    def setup_ui(self):
        """Sets up the user interface."""
        style = ttk.Style()
        style.configure("TLabel", background="#f0f0f0", font=("Arial", 10))
        style.configure("TButton", font=("Arial", 10))
        style.configure("TCombobox", font=("Arial", 10))

        header = tk.Label(self.root, text="Conversor de Arquivos", font=("Arial", 16, "bold"), bg="#f0f0f0")
        header.grid(row=0, column=0, columnspan=3, pady=10)

        tk.Label(self.root, text="Arquivo de Entrada:").grid(row=1, column=0, padx=10, pady=10, sticky="e")
        tk.Entry(self.root, textvariable=self.input_file_path, width=50).grid(row=1, column=1, padx=10, pady=10)
        tk.Button(self.root, text="Selecionar", command=self.select_input_file).grid(row=1, column=2, padx=10, pady=10)

        tk.Label(self.root, text="Tipo de Conversão:").grid(row=2, column=0, padx=10, pady=10, sticky="e")
        conversion_dropdown = ttk.Combobox(self.root, textvariable=self.conversion_type_var, state="readonly")
        conversion_dropdown['values'] = list(FileConverter.conversion_functions.keys())
        conversion_dropdown.grid(row=2, column=1, padx=10, pady=10)

        convert_button = tk.Button(self.root, text="Converter", command=self.perform_conversion_threaded, bg="#4CAF50", fg="white", font=("Arial", 10, "bold"))
        convert_button.grid(row=3, column=1, pady=20)

        progress = ttk.Progressbar(self.root, orient="horizontal", length=300, mode="indeterminate")
        progress.grid(row=4, column=1, pady=10)

    def select_input_file(self):
        """Opens a file dialog to select the input file."""
        file_path = filedialog.askopenfilename(
            title="Selecione o arquivo de entrada",
            filetypes=[("Todos os arquivos", "*.*"), ("PDF Files", "*.pdf"), ("DOCX Files", "*.docx"), ("TXT Files", "*.txt"), ("XLS Files", "*.xls"), ("CSV Files", "*.csv")]
        )
        self.input_file_path.set(file_path)

    def perform_conversion(self):
        """Performs the file conversion."""
        input_file = self.input_file_path.get()
        conversion_type = self.conversion_type_var.get()

        if not input_file:
            messagebox.showerror("Erro", "Por favor, selecione o arquivo de entrada.")
            return

        if not os.path.exists(input_file):
            messagebox.showerror("Erro", "O arquivo de entrada não existe.")
            return

        try:
            conversion_function, output_extension = FileConverter.conversion_functions.get(conversion_type, (None, None))
            if conversion_function:
                base_name, _ = os.path.splitext(input_file)
                output_file = base_name + output_extension
                conversion_function(input_file, output_file)
                logging.info(f"Converted {input_file} to {output_file} ({conversion_type})")
                messagebox.showinfo("Sucesso", f"Arquivo convertido com sucesso! Salvo como: {output_file}")
            else:
                messagebox.showerror("Erro", "Tipo de conversão inválido.")
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro: {str(e)}")

    def perform_conversion_threaded(self):
        """Runs the conversion in a separate thread."""
        threading.Thread(target=self.perform_conversion).start()

    def run(self):
        """Runs the main application loop."""
        self.root.mainloop()


if __name__ == "__main__":
    app = FileConverterApp()
    app.run()